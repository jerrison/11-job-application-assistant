#!/usr/bin/env python3
"""Deterministic resume builder that replicates the Google Doc template exactly.

Takes a JSON file with tailored content and produces a .docx that matches
the source Google Doc formatting pixel-for-pixel.

Usage:
    uv run scripts/build_resume.py input.json -o output/company/resume.docx

Input JSON format:
{
    "tagline": "Principal Product Manager | AI/ML & Enterprise B2B | ...",
    "summary": "Sharp, sufficient 2-3 sentence pitch that communicates the core case...",   // null to omit summary section
    "positions": {
        "moodys": [
            {"bold": "Drove 166% YoY client growth", "text": " for UnderwriteIQ..."},
            ...
        ],
        "kyte": [...],
        "tmobile": [...],
        "lyft": [...],
        "allstate": [...]
    },
    "page_break_before": "tmobile"   // position ID where page 2 starts; null/"auto" allowed
}

All formatting (fonts, sizes, colors, spacing, margins) is hardcoded from the
Google Doc template. The LLM only provides content via the JSON file.
"""

import argparse
import json
import os
import re
import shutil
import subprocess
import tempfile
import time
from pathlib import Path

import json_lenient
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from docx_text import sanitize_docx_text
from lxml import etree
from output_layout import role_submit_dir
from resume_layout import POSITION_ORDER as LAYOUT_POSITION_ORDER
from resume_layout import choose_page_break as layout_choose_page_break

# ─── Constants from Google Doc CSS ─────────────────────────────────────────────
# Extracted from https://docs.google.com/document/d/.../export?format=html

FONT = "Calibri"

# Page layout (540pt content + 36pt padding each side = 612pt = 8.5in)
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(0.5)  # 36pt ≈ 0.5in

# Colors
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_DARK_GRAY = RGBColor(0x44, 0x44, 0x44)
COLOR_SECTION_HEADER = RGBColor(0x2B, 0x57, 0x9A)

# Font sizes (from CSS classes)
SIZE_NAME = Pt(22)  # c23: 22pt bold centered
SIZE_TAGLINE = Pt(10.5)  # c14: 10.5pt, color #444
SIZE_CONTACT = Pt(10)  # c2: 10pt
SIZE_SECTION_HEADER = Pt(14)  # h1: 14pt bold, color #2b579a
SIZE_COMPANY = Pt(11.5)  # c0: 11.5pt bold / c15: 11.5pt normal
SIZE_DATES = Pt(10)  # c2: 10pt italic
SIZE_BULLET = Pt(11)  # c4/c8: 11pt
SIZE_EDU_SCHOOL = Pt(11)  # c4: 11pt bold
SIZE_EDU_DEGREE = Pt(10.5)  # c14: 10.5pt
SIZE_EDU_DETAIL = Pt(10)  # c2: 10pt
SIZE_SKILL_LABEL = Pt(10.5)  # c1: 10.5pt bold
SIZE_SKILL_VALUE = Pt(10.5)  # c14: 10.5pt

# Spacing (from CSS padding values, converted to Pt)
SPACING_LINE = 1.0

# Bullet indent
BULLET_MARGIN_LEFT = Pt(27)
BULLET_PADDING_LEFT = Pt(-9)  # Used via hanging indent
BULLET_CHAR = "\u2022   "  # • followed by 3 nbsp (matching Google Doc: \002022 + 3 spaces)

# ─── Fixed content (never changes) ────────────────────────────────────────────

CANDIDATE_NAME = "JERRISON LI"
CANDIDATE_CONTACT_WITH_LOCATION = (
    "San Francisco, CA  |  jerrisonli@gmail.com  |  510-613-5192  |  linkedin.com/in/jerrison/  |  jerrisonli.com"
)
CANDIDATE_CONTACT_NO_LOCATION = "jerrisonli@gmail.com  |  510-613-5192  |  linkedin.com/in/jerrison/  |  jerrisonli.com"

POSITIONS_META = {
    "moodys": {
        "company": "MOODY'S ANALYTICS",
        "title": "Associate Director, Product Management",
        "dates": "San Francisco, CA | 2024\u2013Present",
    },
    "kyte": {
        "company": "KYTE",
        "title": "Staff Product Manager",
        "subtitle": " (Series B, On-Demand Car Rental)",
        "dates": "San Francisco, CA | 2022\u20132024 | 150 employees, 15K+ monthly transactions",
    },
    "tmobile": {
        "company": "T-MOBILE",
        "title": "Senior Product Manager, IoT Platform",
        "dates": "New York, NY | 2020\u20132022 | Enterprise IoT Business Unit",
    },
    "lyft": {
        "company": "LYFT",
        "title": "Senior Actuarial Analyst (First Actuary Hire)",
        "dates": "San Francisco, CA | 2016\u20132017",
    },
    "allstate": {
        "company": "ALLSTATE",
        "title": "Senior Actuarial Analyst",
        "dates": "San Francisco, CA | 2013\u20132016",
    },
}

POSITION_ORDER = ["moodys", "kyte", "tmobile", "lyft", "allstate"]

# Approximate layout model used to auto-balance page 1 vs page 2.
# This is intentionally conservative; validation still checks the final PDF.
USABLE_PAGE_HEIGHT_PT = 720
PAGE_BREAK_SAFETY_BUFFER_PT = 12
TITLE_CHARS_PER_LINE = 84
SUMMARY_CHARS_PER_LINE = 104
BULLET_CHARS_PER_LINE = 92
DETAIL_CHARS_PER_LINE = 98
SKILL_CHARS_PER_LINE = 96
CONTACT_CHARS_PER_LINE = 100

CALIFORNIA_LOCATION_RE = re.compile(r"(?:,\s*CA\b|\bCalifornia\b)", re.I)
CALIFORNIA_CITY_TOKENS = (
    "berkeley",
    "cupertino",
    "fremont",
    "irvine",
    "los angeles",
    "menlo park",
    "mountain view",
    "oakland",
    "palo alto",
    "pasadena",
    "redwood city",
    "sacramento",
    "san diego",
    "san francisco",
    "san jose",
    "santa clara",
    "santa monica",
    "sunnyvale",
)
LOCATION_LINE_RE = re.compile(r"(?:\*\*)?Location(?:s)?(?:\*\*)?:?\s*(.+)", re.I)
JSON_STRING_RE = re.compile(r'"((?:[^"\\]|\\.)*)"')
STRUCTURED_LOCATION_KEYS = (
    "locationName",
    "locationExternalName",
    "secondaryLocationNames",
)


def _unique_ordered(values: list[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for value in values:
        clean = value.strip()
        if not clean:
            continue
        key = clean.casefold()
        if key in seen:
            continue
        seen.add(key)
        ordered.append(clean)
    return ordered


def _extract_location_hints_from_raw_text(text: str) -> list[str]:
    hints: list[str] = []
    for line in text.splitlines()[:20]:
        match = LOCATION_LINE_RE.search(line.strip())
        if not match:
            continue
        value = match.group(1).strip().strip("*").strip()
        if value:
            hints.append(value)
    return _unique_ordered(hints)


def _extract_location_hints_from_application_html(html_text: str) -> list[str]:
    hints: list[str] = []

    for key in ("locationName", "locationExternalName"):
        pattern = re.compile(rf'"{re.escape(key)}"\s*:\s*"((?:[^"\\]|\\.)*)"', re.I)
        for match in pattern.finditer(html_text):
            hints.append(json.loads(f'"{match.group(1)}"'))

    secondary_match = re.search(r'"secondaryLocationNames"\s*:\s*\[(.*?)\]', html_text, re.I | re.S)
    if secondary_match:
        for raw in JSON_STRING_RE.findall(secondary_match.group(1)):
            hints.append(json.loads(f'"{raw}"'))

    return _unique_ordered(hints)


def _load_job_location(input_path: str | os.PathLike[str] | None) -> str | None:
    """Load location hints from resume-adjacent JD artifacts when available."""
    if not input_path:
        return None

    root = Path(input_path).resolve().parent
    role_root = root.parent if (root.parent / ".pipeline_meta.json").exists() else root
    hints: list[str] = []

    parsed_path = root / "jd_parsed.json"
    if parsed_path.exists():
        try:
            data = json_lenient.loads(parsed_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            data = {}
        location = str(data.get("location", "")).strip()
        if location:
            hints.append(location)

    raw_path = root / "jd_raw.md"
    if raw_path.exists():
        try:
            hints.extend(_extract_location_hints_from_raw_text(raw_path.read_text(encoding="utf-8")))
        except OSError:
            pass

    submit_dir = role_submit_dir(role_root)
    html_candidates = sorted(submit_dir.glob("*_application_page.html")) if submit_dir.exists() else []
    if not html_candidates:
        html_candidates = sorted(root.glob("*_application_page.html"))

    for html_path in html_candidates:
        try:
            hints.extend(_extract_location_hints_from_application_html(html_path.read_text(encoding="utf-8")))
        except OSError:
            continue

    unique_hints = _unique_ordered(hints)
    if not unique_hints:
        return None
    return " | ".join(unique_hints)


def _is_california_role(job_location: str | None) -> bool:
    if not job_location:
        return False

    if CALIFORNIA_LOCATION_RE.search(job_location):
        return True

    lowered = job_location.casefold()
    return any(token in lowered for token in CALIFORNIA_CITY_TOKENS)


def _candidate_contact_for_input(input_path: str | os.PathLike[str] | None) -> str:
    """Use the city-prefixed contact line only for California roles."""
    job_location = _load_job_location(input_path)
    if job_location is None:
        return CANDIDATE_CONTACT_WITH_LOCATION
    if _is_california_role(job_location):
        return CANDIDATE_CONTACT_WITH_LOCATION
    return CANDIDATE_CONTACT_NO_LOCATION


def _estimate_wrapped_lines(text: str, chars_per_line: int) -> int:
    """Roughly estimate wrapped lines for plain text."""
    if not text:
        return 0

    text = text.replace("\u2014", "-")
    paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
    if not paragraphs:
        return 0

    lines = 0
    for paragraph in paragraphs:
        words = paragraph.split()
        if not words:
            continue

        current = 0
        paragraph_lines = 1
        for word in words:
            word_len = len(word)
            if current == 0:
                current = word_len
                continue
            if current + 1 + word_len <= chars_per_line:
                current += 1 + word_len
            else:
                paragraph_lines += 1
                current = word_len
        lines += paragraph_lines

    return max(lines, 1)


def _estimate_position_height(meta: dict, bullets: list) -> float:
    """Estimate the rendered height of one position block in points."""
    title_text = meta["company"] + " - " + meta["title"] + meta.get("subtitle", "")
    title_lines = _estimate_wrapped_lines(title_text, TITLE_CHARS_PER_LINE)
    dates_lines = _estimate_wrapped_lines(meta["dates"], TITLE_CHARS_PER_LINE)

    # before 3 + after 2 + ~11.5pt per wrapped line
    height = 5 + (12 * title_lines)
    # before 0 + after 4 + ~10pt per wrapped line
    height += 4 + (10 * dates_lines)

    for bullet in bullets:
        if isinstance(bullet, dict):
            bullet_text = (bullet.get("bold", "") + bullet.get("text", "")).strip()
        else:
            bullet_text = str(bullet).strip()
        bullet_lines = _estimate_wrapped_lines(bullet_text, BULLET_CHARS_PER_LINE)
        # before 0 + after 3 + ~11pt per wrapped line
        height += 3 + (11 * bullet_lines)

    return height


def _estimate_summary_height(summary: str | None) -> float:
    if not summary:
        return 0
    summary_lines = _estimate_wrapped_lines(summary, SUMMARY_CHARS_PER_LINE)
    # section header + summary paragraph
    return 30 + 6 + (11 * summary_lines)


def _estimate_contact_height(contact_line: str) -> float:
    contact_lines = _estimate_wrapped_lines(contact_line, CONTACT_CHARS_PER_LINE)
    return 10 + (10 * contact_lines)


def _estimate_static_page1_height(data: dict, contact_line: str) -> float:
    # name + tagline + contact + experience header + optional summary block
    return 25 + 15 + _estimate_contact_height(contact_line) + _estimate_summary_height(data.get("summary")) + 30


def _estimate_education_skills_height() -> float:
    height = 30  # EDUCATION header

    for edu in EDUCATION:
        school_lines = _estimate_wrapped_lines(edu["school"], TITLE_CHARS_PER_LINE)
        degree_text = edu["degree"] + " | " + edu["location"]
        degree_lines = _estimate_wrapped_lines(degree_text, DETAIL_CHARS_PER_LINE)
        height += 2 + (11 * school_lines)
        height += 2 + (11 * degree_lines)
        for detail in edu["details"]:
            detail_lines = _estimate_wrapped_lines(detail, DETAIL_CHARS_PER_LINE)
            height += 5 + (10 * detail_lines)

    height += 30  # SKILLS header
    for label, value in SKILLS:
        skill_lines = _estimate_wrapped_lines(label + " " + value, SKILL_CHARS_PER_LINE)
        height += 3 + (10.5 * skill_lines)

    return height


def choose_page_break(data: dict, contact_line: str | None = None) -> tuple[str | None, dict]:
    """Choose the latest page break that keeps page 1 dense without overflowing.

    Returns (page_break_before, diagnostics).
    """
    included_positions = [pos_id for pos_id in POSITION_ORDER if data.get("positions", {}).get(pos_id)]

    diagnostics = {
        "included_positions": included_positions,
        "candidates": [],
    }

    if len(included_positions) <= 1:
        diagnostics["reason"] = "Only one populated position block."
        return data.get("page_break_before"), diagnostics

    static_page1 = _estimate_static_page1_height(data, contact_line or CANDIDATE_CONTACT_WITH_LOCATION)
    static_page2 = _estimate_education_skills_height()
    position_heights = {
        pos_id: _estimate_position_height(POSITIONS_META[pos_id], data["positions"][pos_id])
        for pos_id in included_positions
    }

    best_candidate = None
    overflow_candidate = None

    for index, candidate in enumerate(included_positions[1:], start=1):
        page1_ids = included_positions[:index]
        page2_ids = included_positions[index:]
        page1_height = static_page1 + sum(position_heights[pos_id] for pos_id in page1_ids)
        page2_height = static_page2 + sum(position_heights[pos_id] for pos_id in page2_ids)

        candidate_diag = {
            "page_break_before": candidate,
            "page1_positions": page1_ids,
            "page2_positions": page2_ids,
            "page1_height": round(page1_height, 1),
            "page2_height": round(page2_height, 1),
        }
        diagnostics["candidates"].append(candidate_diag)

        if page1_height <= (USABLE_PAGE_HEIGHT_PT - PAGE_BREAK_SAFETY_BUFFER_PT):
            best_candidate = candidate_diag
        elif overflow_candidate is None:
            overflow_candidate = candidate_diag

    if best_candidate is not None:
        diagnostics["selected"] = best_candidate
        diagnostics["reason"] = "Latest page break that keeps page 1 comfortably within one page."
        return best_candidate["page_break_before"], diagnostics

    diagnostics["selected"] = overflow_candidate
    diagnostics["reason"] = "No candidate kept page 1 within one page estimate; chose earliest break."
    return overflow_candidate["page_break_before"], diagnostics


EDUCATION = [
    {
        "school": "THE WHARTON SCHOOL, UNIVERSITY OF PENNSYLVANIA",
        "degree": "MBA (Finance)",
        "location": "Philadelphia, PA | 2018\u20132020",
        "details": [
            "Joseph Wharton Fellow (academic, personal, and professional achievement)",
            "GMAT 750 (98th percentile)",
        ],
    },
    {
        "school": "PENN ENGINEERING, UNIVERSITY OF PENNSYLVANIA",
        "degree": "M.S. Computer Science",
        "location": "Philadelphia, PA | 2018\u20132020",
        "details": [
            "Ripple Research Fellow (blockchain applications in P&C insurance). Relevant coursework: Applied ML, Database Systems, AI.",
        ],
    },
    {
        "school": "FLORIDA STATE UNIVERSITY",
        "degree": "B.S. Actuarial Science & Computational Science (Dual Degree)",
        "location": "Tallahassee, FL | 2009\u20132013",
        "details": ["Phi Beta Kappa, Magna Cum Laude"],
        "details_italic": True,
    },
]

SKILLS = [
    ("Competitions:", "Gold Medalist, National Physics Olympiad, Panama (3,500 students)"),
    (
        "Technical:",
        "Python, SQL, TypeScript, Figma | ML/AI: Snowflake Cortex, LLM orchestration, RAG systems, GLMs | Data: A/B testing, analytics pipelines",
    ),
    ("Languages:", "Spanish (native), Cantonese (native), Mandarin (advanced)"),
    ("Certifications:", "Associate of the Casualty Actuarial Society (ACAS)"),
    ("Work Authorization:", "United States Citizen"),
]


# ─── Builder ──────────────────────────────────────────────────────────────────


def _run(para, text, size, bold=False, italic=False, color=COLOR_BLACK, font=FONT):
    """Add a run to a paragraph with exact formatting."""
    run = para.add_run(sanitize_docx_text(text))
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # Force font via XML (required for some renderers)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rFonts.set(qn("w:cs"), font)
    return run


def _set_spacing(para, before=None, after=None, line=SPACING_LINE):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    if before is not None:
        pf.space_before = before
    if after is not None:
        pf.space_after = after
    pf.line_spacing = line


def _add_bottom_border(para, color="2B579A", size="6"):
    """Add a bottom border to a paragraph (for section headers)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)  # 6 eighths of a point ≈ 0.75pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def build_resume(
    data: dict,
    output_path: str,
    dry_run: bool = False,
    input_path: str | os.PathLike[str] | None = None,
):
    """Build the resume .docx from structured content data."""
    doc = Document()
    contact_line = _candidate_contact_for_input(input_path)

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # Set default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = SIZE_BULLET
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = SPACING_LINE

    # ── Name (22pt Calibri bold, centered) ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=Pt(0), after=Pt(3))
    _run(p, CANDIDATE_NAME, SIZE_NAME, bold=True)

    # ── Tagline (10.5pt, color #444, centered) ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=Pt(0), after=Pt(4))
    _run(p, data["tagline"], SIZE_TAGLINE, color=COLOR_DARK_GRAY)

    # ── Contact (10pt, centered, 10pt bottom padding) ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=Pt(0), after=Pt(10))
    _run(p, contact_line, SIZE_CONTACT)

    # ── Summary section ──
    if data.get("summary"):
        _add_section_header(doc, "Summary")
        p = doc.add_paragraph()
        _set_spacing(p, before=Pt(0), after=Pt(6))
        _run(p, data["summary"], SIZE_BULLET)

    # ── Experience section ──
    _add_section_header(doc, "EXPERIENCE")

    page_break_before = data.get("page_break_before")
    if page_break_before in (None, "", "auto"):
        page_break_before, diagnostics = layout_choose_page_break(data)
        data["page_break_before"] = page_break_before
        if page_break_before:
            print(f"Auto-selected page_break_before: {page_break_before}")
            if diagnostics.get("reason"):
                print(f"  {diagnostics['reason']}")

    for pos_id in LAYOUT_POSITION_ORDER:
        bullets = data["positions"].get(pos_id, [])
        if not bullets:
            continue

        meta = POSITIONS_META[pos_id]

        # Page break
        if pos_id == page_break_before:
            # Insert page break on the company name paragraph
            _add_position(doc, meta, bullets, page_break=True)
        else:
            _add_position(doc, meta, bullets, page_break=False)

    # ── Education section ──
    _add_section_header(doc, "EDUCATION")
    for edu in EDUCATION:
        # School name (11pt bold)
        p = doc.add_paragraph()
        _set_spacing(p, before=Pt(0), after=Pt(2))
        _run(p, edu["school"], SIZE_EDU_SCHOOL, bold=True)

        # Degree + location (10.5pt + 10pt italic)
        p = doc.add_paragraph()
        _set_spacing(p, before=Pt(0), after=Pt(2))
        _run(p, edu["degree"], SIZE_EDU_DEGREE)
        _run(p, " | " + edu["location"], SIZE_DATES, italic=True)

        # Details
        is_italic = edu.get("details_italic", False)
        for detail in edu["details"]:
            p = doc.add_paragraph()
            _set_spacing(p, before=Pt(0), after=Pt(5))
            _run(p, detail, SIZE_EDU_DETAIL, italic=is_italic)

    # ── Skills section ──
    _add_section_header(doc, "SKILLS & ADDITIONAL")
    for label, value in SKILLS:
        p = doc.add_paragraph()
        _set_spacing(p, before=Pt(0), after=Pt(3))
        _run(p, label + " ", SIZE_SKILL_LABEL, bold=True)
        _run(p, value, SIZE_SKILL_VALUE)

    # ── Save ──
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"Resume saved: {output_path}")

    # ── Convert to PDF via LibreOffice ──
    if dry_run:
        print("Dry run: skipping PDF conversion")
        return

    _convert_to_pdf(output_path)

    # ── Auto-validate PDF ──
    base = os.path.splitext(os.path.basename(output_path))[0]
    pdf_path = os.path.join(os.path.dirname(os.path.abspath(output_path)) or ".", base + ".pdf")
    if os.path.exists(pdf_path):
        _validate_pdf(pdf_path)


def _convert_to_pdf(docx_path: str):
    """Convert .docx to .pdf using LibreOffice."""
    soffice = _find_libreoffice()
    if not soffice:
        print("Warning: LibreOffice not found — skipping PDF conversion.")
        print("Install with: brew install --cask libreoffice")
        return

    out_dir = os.path.dirname(os.path.abspath(docx_path)) or "."
    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(out_dir, base + ".pdf")
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    env = os.environ.copy()
    # Force LibreOffice onto a non-AppKit VCL backend. On current macOS builds,
    # the default osx plugin can abort even for `--headless` conversions. This
    # must override inherited shell state such as SAL_USE_VCLPLUGIN=osx.
    env["SAL_USE_VCLPLUGIN"] = "svp"
    with tempfile.TemporaryDirectory(prefix="libreoffice-profile-") as profile_dir:
        result = subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={Path(profile_dir).resolve().as_uri()}",
                "--headless",
                "--nologo",
                "--nodefault",
                "--nolockcheck",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                out_dir,
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
            env=env,
        )

    deadline = time.monotonic() + 5
    while time.monotonic() < deadline and not os.path.exists(pdf_path):
        time.sleep(0.2)
    if os.path.exists(pdf_path):
        print(f"Resume PDF saved: {pdf_path}")
        if result.returncode != 0:
            detail = (result.stderr or result.stdout or "").strip()
            if detail:
                print(f"Note: LibreOffice returned a non-zero exit code after producing the PDF: {detail}")
        return

    detail = (result.stderr or result.stdout or "").strip()
    if result.returncode == 0:
        print("Warning: PDF conversion ran but output file not found.")
    else:
        print(f"Warning: PDF conversion failed: {detail}")


def _validate_pdf(pdf_path: str) -> bool:
    """Validate the built PDF (page count, name presence)."""
    try:
        import pdfplumber
    except ImportError:
        print("Warning: pdfplumber not installed — skipping PDF validation.")
        return False

    print(f"\n── Validating {pdf_path} ──")
    all_passed = True

    with pdfplumber.open(pdf_path) as pdf:
        page_count = len(pdf.pages)
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    # Check page count == 2
    if page_count == 2:
        print(f"PASS: Page count is {page_count}")
    else:
        print(f"FAIL: Page count is {page_count} (expected 2)")
        all_passed = False

    # Check candidate name in text
    if "JERRISON LI" in text:
        print("PASS: 'JERRISON LI' found in PDF text")
    else:
        print("FAIL: 'JERRISON LI' not found in PDF text")
        all_passed = False

    if not all_passed:
        print("\nRemediation guidance:")
        if page_count != 2:
            print("  - Adjust bullet content, then re-run scripts/optimize_page_break.py before rebuilding.")
        if "JERRISON LI" not in text:
            print("  - Check that the name renders correctly; font embedding may be broken.")

    return all_passed


def _find_libreoffice() -> str | None:
    """Find the LibreOffice soffice binary."""
    # Check PATH first
    soffice = shutil.which("soffice")
    if soffice:
        return soffice
    # macOS default install location
    mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.exists(mac_path):
        return mac_path
    return None


def _add_section_header(doc, title):
    """Add a section header with bottom border (14pt bold #2b579a)."""
    p = doc.add_paragraph()
    _set_spacing(p, before=Pt(10), after=Pt(6))
    _run(p, title, SIZE_SECTION_HEADER, bold=True, color=COLOR_SECTION_HEADER)
    _add_bottom_border(p)


def _add_position(doc, meta, bullets, page_break=False):
    """Add a position block (company + title + dates + bullets)."""
    # Company + Title line (11.5pt; company bold, title normal)
    p = doc.add_paragraph()
    _set_spacing(p, before=Pt(3), after=Pt(2))
    if page_break:
        p.paragraph_format.page_break_before = True
    _run(p, meta["company"], SIZE_COMPANY, bold=True)
    _run(p, " \u2014 " + meta["title"], SIZE_COMPANY)
    if "subtitle" in meta:
        _run(p, meta["subtitle"], SIZE_DATES, italic=True)

    # Dates line (10pt italic)
    p = doc.add_paragraph()
    _set_spacing(p, before=Pt(0), after=Pt(4))
    _run(p, meta["dates"], SIZE_DATES, italic=True)

    # Bullets
    for bullet in bullets:
        p = doc.add_paragraph()
        _set_spacing(p, before=Pt(0), after=Pt(3))
        p.paragraph_format.left_indent = BULLET_MARGIN_LEFT
        p.paragraph_format.first_line_indent = Pt(-18)  # Hanging indent for bullet char

        # Bullet character
        _run(p, BULLET_CHAR, SIZE_BULLET)

        # Bold lead + normal body
        if isinstance(bullet, dict):
            _run(p, bullet["bold"], SIZE_BULLET, bold=True)
            _run(p, bullet["text"], SIZE_BULLET)
        else:
            # Plain string bullet (no bold/normal split)
            _run(p, bullet, SIZE_BULLET)


def main():
    parser = argparse.ArgumentParser(description="Build resume .docx from JSON content")
    parser.add_argument("input", help="Path to JSON content file")
    parser.add_argument("-o", "--output", default="output/resume.docx", help="Output .docx path")
    parser.add_argument("--dry-run", action="store_true", help="Build .docx but skip PDF conversion")
    args = parser.parse_args()

    with open(args.input) as f:
        data = json_lenient.load(f)

    build_resume(data, args.output, dry_run=args.dry_run, input_path=args.input)


if __name__ == "__main__":
    main()
