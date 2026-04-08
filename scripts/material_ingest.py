#!/usr/bin/env python3
"""Helpers for importing user-provided source materials."""

from __future__ import annotations

import io
import re
from pathlib import Path
from urllib.parse import urlsplit

import httpx
from docx import Document
from lxml import html
from pypdf import PdfReader

_GOOGLE_DOC_RE = re.compile(r"^https?://docs\.google\.com/document/d/([A-Za-z0-9_-]+)")


class FetchResult:
    def __init__(self, *, url: str, content_type: str, content: bytes):
        self.url = url
        self.content_type = content_type
        self.content = content


def normalize_source_url(source_url: str) -> str:
    normalized = str(source_url or "").strip()
    if not normalized:
        raise ValueError("source_url is required")

    match = _GOOGLE_DOC_RE.match(normalized)
    if match:
        doc_id = match.group(1)
        return f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
    return normalized


def fetch_source_url(source_url: str) -> FetchResult:
    normalized = normalize_source_url(source_url)
    response = httpx.get(
        normalized,
        follow_redirects=True,
        timeout=20.0,
        headers={
            "User-Agent": "job-assets/1.0",
            "Accept": (
                "text/plain, text/markdown, text/html, application/pdf, "
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document;q=0.9, */*;q=0.8"
            ),
        },
    )
    response.raise_for_status()
    return FetchResult(
        url=str(response.url),
        content_type=response.headers.get("content-type", ""),
        content=response.content,
    )


def _normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _finalize_extracted_text(text: str, *, preserve_trailing_newline: bool = False) -> str:
    normalized = _normalize_newlines(text)
    if preserve_trailing_newline:
        return normalized
    stripped = normalized.strip()
    if not stripped:
        raise ValueError("No importable text found in source material")
    return f"{stripped}\n"


def _extract_docx_text(content_bytes: bytes) -> str:
    document = Document(io.BytesIO(content_bytes))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            chunks.append(value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    chunks.append(value)
    return _finalize_extracted_text("\n\n".join(chunks))


def _extract_pdf_text(content_bytes: bytes) -> str:
    texts: list[str] = []
    reader = PdfReader(io.BytesIO(content_bytes))
    for page in reader.pages:
        extracted = (page.extract_text() or "").strip()
        if extracted:
            texts.append(extracted)

    if not texts:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            for page in pdf.pages:
                extracted = (page.extract_text() or "").strip()
                if extracted:
                    texts.append(extracted)

    return _finalize_extracted_text("\n\n".join(texts))


def _extract_html_text(content_bytes: bytes) -> str:
    document = html.fromstring(content_bytes)
    for node in document.xpath("//script|//style|//noscript"):
        node.drop_tree()
    text = document.text_content()
    return _finalize_extracted_text(text)


def _extract_text_bytes(content_bytes: bytes) -> str:
    try:
        decoded = content_bytes.decode("utf-8")
    except UnicodeDecodeError:
        decoded = content_bytes.decode("utf-8", errors="replace")
    return _finalize_extracted_text(decoded, preserve_trailing_newline=True)


def _content_kind(*, file_name: str | None, content_type: str | None) -> str:
    lower_name = str(file_name or "").casefold()
    lower_type = str(content_type or "").casefold()
    if lower_name.endswith(".docx") or "officedocument.wordprocessingml.document" in lower_type:
        return "docx"
    if lower_name.endswith(".pdf") or "application/pdf" in lower_type:
        return "pdf"
    if lower_name.endswith((".html", ".htm")) or "text/html" in lower_type:
        return "html"
    return "text"


def _guess_file_name_from_url(source_url: str) -> str:
    path = urlsplit(source_url).path
    return Path(path).name


def import_material_content(
    *,
    text: str | None = None,
    source_url: str | None = None,
    file_name: str | None = None,
    content_type: str | None = None,
    content_bytes: bytes | None = None,
    fetcher=None,
) -> str:
    has_text = text is not None and text != ""
    has_url = bool(str(source_url or "").strip())
    has_bytes = content_bytes is not None
    provided_count = int(has_text) + int(has_url) + int(has_bytes)
    if provided_count != 1:
        raise ValueError("Provide exactly one material source: text, source_url, or content_bytes")

    if has_text:
        return _finalize_extracted_text(str(text), preserve_trailing_newline=True)

    if has_url:
        fetch = fetcher or fetch_source_url
        normalized_source_url = normalize_source_url(str(source_url).strip())
        result = fetch(normalized_source_url)
        return import_material_content(
            file_name=_guess_file_name_from_url(result.url),
            content_type=result.content_type,
            content_bytes=result.content,
        )

    kind = _content_kind(file_name=file_name, content_type=content_type)
    if kind == "docx":
        return _extract_docx_text(content_bytes or b"")
    if kind == "pdf":
        return _extract_pdf_text(content_bytes or b"")
    if kind == "html":
        return _extract_html_text(content_bytes or b"")
    return _extract_text_bytes(content_bytes or b"")
