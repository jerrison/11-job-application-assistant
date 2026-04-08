#!/usr/bin/env python3
"""Generate a personalized interview preparation guide.

Gathers context from the job's output directory (JD, research, resume, work stories),
spawns the configured provider subprocess to generate a comprehensive prep guide in markdown,
then converts it to .docx and .pdf.

Usage:
    uv run python scripts/generate_interview_prep.py output/uber/senior-pm-web --force
    uv run python scripts/generate_interview_prep.py output/uber/senior-pm-web --stage "Phone Screen"
    uv run python scripts/generate_interview_prep.py output/uber/senior-pm-web --interviewer "Jane Doe" --interviewer "John Smith"
"""

from __future__ import annotations

import argparse
import datetime as _dt
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from docx_text import sanitize_docx_text

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from llm_provider import default_active_provider, provider_available, provider_command

PROJECT_ROOT = Path(__file__).resolve().parent.parent
log = logging.getLogger(__name__)

FONT = "Calibri"
COLOR_BLACK = RGBColor(0x33, 0x33, 0x33)
COLOR_BLUE = RGBColor(0x2B, 0x57, 0x9A)
COLOR_GRAY = RGBColor(0x55, 0x55, 0x55)
CANDIDATE_NAME = "Jerrison Li"


# ---------------------------------------------------------------------------
# Progress tracking
# ---------------------------------------------------------------------------


def _write_progress(prep_dir: Path, status: str, detail: str) -> None:
    prep_dir.mkdir(parents=True, exist_ok=True)
    (prep_dir / ".progress.json").write_text(
        json.dumps(
            {
                "status": status,
                "detail": detail,
                "updated_at": datetime.now(_dt.UTC).isoformat(),
            }
        )
    )


# ---------------------------------------------------------------------------
# Context gathering
# ---------------------------------------------------------------------------


def _read_optional(path: Path) -> str:
    """Read file contents, returning empty string if missing."""
    if path.exists():
        return path.read_text(encoding="utf-8")
    return ""


def _gather_context(out_dir: Path) -> dict:
    """Read all context files and return a dict with their contents."""
    content_dir = out_dir / "content"
    meta_path = out_dir / ".pipeline_meta.json"

    meta = {}
    meta_text = _read_optional(meta_path)
    if meta_text:
        meta = json.loads(meta_text)

    return {
        "meta": meta,
        "company": meta.get("company_proper", meta.get("company", "Unknown")),
        "jd_title": meta.get("jd_title", "Unknown Role"),
        "jd_source": meta.get("jd_source", ""),
        "jd_raw": _read_optional(content_dir / "jd_raw.md"),
        "jd_parsed": _read_optional(content_dir / "jd_parsed.json"),
        "research_cache": _read_optional(content_dir / "role_research_cache.json"),
        "master_resume": _read_optional(PROJECT_ROOT / "master_resume.md"),
        "work_stories": _read_optional(PROJECT_ROOT / "work_stories.md"),
        "candidate_context": _read_optional(PROJECT_ROOT / "candidate_context.md"),
        "application_profile": _read_optional(PROJECT_ROOT / "application_profile.md"),
    }


# ---------------------------------------------------------------------------
# Prompt construction
# ---------------------------------------------------------------------------


def _build_prompt(
    ctx: dict,
    prep_dir: Path,
    stage: str,
    interviewers: list[str],
    notes: str,
) -> str:
    """Build the full prompt (system + user message) for the configured provider."""
    system_prompt_path = PROJECT_ROOT / "scripts" / "prompts" / "interview_prep_system.md"
    system_prompt = system_prompt_path.read_text(encoding="utf-8")

    output_path = prep_dir / "interview_prep.md"

    # Build the user message with all context sections
    user_parts = []
    user_parts.append(
        f"Generate a comprehensive interview preparation guide for the role of "
        f"**{ctx['jd_title']}** at **{ctx['company']}**."
    )
    user_parts.append(f"\n**Interview stage:** {stage}")

    if interviewers:
        interviewer_list = ", ".join(interviewers)
        user_parts.append(f"**Interviewers:** {interviewer_list}")

    if notes:
        user_parts.append(f"**Additional notes:** {notes}")

    if ctx["jd_source"]:
        user_parts.append(f"**Job posting URL:** {ctx['jd_source']}")

    user_parts.append(f"\n**Write the complete guide to:** `{output_path}`")

    # Context sections
    user_parts.append(f"\n<job_description_raw>\n{ctx['jd_raw']}\n</job_description_raw>")
    user_parts.append(f"\n<job_description_parsed>\n{ctx['jd_parsed']}\n</job_description_parsed>")
    user_parts.append(f"\n<research_cache>\n{ctx['research_cache']}\n</research_cache>")
    user_parts.append(f"\n<master_resume>\n{ctx['master_resume']}\n</master_resume>")
    user_parts.append(f"\n<work_stories>\n{ctx['work_stories']}\n</work_stories>")
    user_parts.append(f"\n<candidate_context>\n{ctx['candidate_context']}\n</candidate_context>")
    user_parts.append(f"\n<application_profile>\n{ctx['application_profile']}\n</application_profile>")

    user_message = "\n".join(user_parts)
    full_prompt = f"<system>\n{system_prompt}\n</system>\n\n{user_message}"
    return full_prompt


# ---------------------------------------------------------------------------
# Provider subprocess
# ---------------------------------------------------------------------------


def _run_provider(full_prompt: str, prep_dir: Path) -> str:
    """Spawn the configured provider and return the generated markdown content."""
    provider = default_active_provider()
    if not provider_available(provider):
        log.warning("Interview prep provider %s is not available — skipping", provider)
        return ""

    output_path = prep_dir / "interview_prep.md"

    claude_allowed_tools = ",".join(
        [
            "WebSearch",
            "WebFetch",
            "Read",
            "Write",
            "Glob",
            "Grep",
            "mcp__plugin_playwright_playwright__browser_navigate",
            "mcp__plugin_playwright_playwright__browser_snapshot",
            "mcp__plugin_playwright_playwright__browser_click",
            "mcp__plugin_playwright_playwright__browser_take_screenshot",
        ]
    )
    allowed_tools = claude_allowed_tools if provider == "claude" else None
    command = provider_command(
        provider,
        full_prompt,
        project_root=PROJECT_ROOT,
        search_enabled=True,
        file_tools_enabled=(provider == "openai"),
        claude_allowed_tools=allowed_tools,
    )

    result = subprocess.run(
        command,
        cwd=str(PROJECT_ROOT),
        capture_output=True,
        text=True,
        timeout=900,
    )

    if result.returncode != 0:
        stderr_snippet = (result.stderr or "")[:500]
        raise RuntimeError(f"{provider} provider exited with code {result.returncode}: {stderr_snippet}")

    # Prefer file output when the provider used write/file tools.
    if output_path.exists():
        return output_path.read_text(encoding="utf-8")

    # Fallback: save stdout as the markdown file
    content = result.stdout.strip()
    if not content:
        raise RuntimeError(f"{provider} provider produced no output and did not write the file.")

    prep_dir.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content, encoding="utf-8")
    return content


# ---------------------------------------------------------------------------
# Document generation: markdown -> .docx
# ---------------------------------------------------------------------------


def _docx_run(para, text, size, bold=False, italic=False, color=COLOR_BLACK):
    """Add a styled run to a paragraph with cross-platform font compatibility."""
    run = para.add_run(sanitize_docx_text(text))
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>')
        rPr.insert(0, rFonts)
    return run


def _apply_inline_formatting(para, text, size, base_bold=False, base_italic=False, color=COLOR_BLACK):
    """Parse inline **bold** and *italic* markers and add runs accordingly."""
    # Pattern: **bold**, *italic*, or plain text segments
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))")
    for m in pattern.finditer(text):
        if m.group(2) is not None:
            # **bold**
            _docx_run(para, m.group(2), size, bold=True, italic=base_italic, color=color)
        elif m.group(3) is not None:
            # *italic*
            _docx_run(para, m.group(3), size, bold=base_bold, italic=True, color=color)
        else:
            # plain text
            _docx_run(para, m.group(4), size, bold=base_bold, italic=base_italic, color=color)


def _build_docx(md_text: str, docx_path: Path) -> None:
    """Convert markdown text to a styled .docx document."""
    doc = Document()

    # Page setup: US Letter, 0.75in margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    in_code_block = False
    lines = md_text.split("\n")

    for line in lines:
        stripped = line.strip()

        # --- horizontal rule: skip
        if stripped in ("---", "***", "___"):
            continue

        # Code block toggle
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        # Inside code block
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_BLACK
            # Set Courier New in XML for cross-platform
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(
                    f'<w:rFonts {nsdecls("w")} w:ascii="Courier New" w:hAnsi="Courier New" w:cs="Courier New"/>'
                )
                rPr.insert(0, rFonts)
            continue

        # Headings
        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            _docx_run(p, stripped[5:], Pt(11), bold=True, italic=True)
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            _docx_run(p, stripped[4:], Pt(12), bold=True)
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            _docx_run(p, stripped[3:], Pt(14), bold=True, color=COLOR_BLUE)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(8)
            _docx_run(p, stripped[2:], Pt(18), bold=True, color=COLOR_BLUE)
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            _apply_inline_formatting(p, stripped[2:], Pt(11), base_italic=True, color=COLOR_GRAY)
            continue

        # Bullet points (- or *)
        if stripped.startswith(("- ", "* ")):
            bullet_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(2)

            # Add bullet character
            _docx_run(p, "\u2022  ", Pt(11))

            # Handle bold prefix pattern: **text**: detail or **text** detail
            bold_prefix_match = re.match(r"^\*\*(.+?)\*\*(:\s*|\s+)(.*)", bullet_text)
            if bold_prefix_match:
                _docx_run(p, bold_prefix_match.group(1), Pt(11), bold=True)
                separator = bold_prefix_match.group(2)
                rest = bold_prefix_match.group(3)
                if rest:
                    _docx_run(p, separator if separator.startswith(":") else " ", Pt(11))
                    _apply_inline_formatting(p, rest, Pt(11))
            else:
                _apply_inline_formatting(p, bullet_text, Pt(11))
            continue

        # Nested bullet (  - or  *)
        if re.match(r"^\s{2,}[-*] ", line):
            bullet_text = re.sub(r"^\s{2,}[-*] ", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(2)
            _docx_run(p, "\u2013  ", Pt(10))  # en-dash for sub-bullets
            _apply_inline_formatting(p, bullet_text, Pt(10))
            continue

        # Empty line
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            continue

        # Regular text
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _apply_inline_formatting(p, stripped, Pt(11))

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Interview prep .docx saved: {docx_path}")


# ---------------------------------------------------------------------------
# Document generation: .docx -> .pdf
# ---------------------------------------------------------------------------


def _find_libreoffice() -> str | None:
    """Find the LibreOffice soffice binary."""
    mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.exists(mac_path):
        return mac_path
    soffice = shutil.which("soffice")
    if soffice:
        return soffice
    return None


def _convert_to_pdf(docx_path: Path) -> None:
    """Convert .docx to .pdf using LibreOffice headless."""
    soffice = _find_libreoffice()
    if not soffice:
        print("Warning: LibreOffice not found — skipping PDF conversion.")
        print("Install with: brew install --cask libreoffice")
        return

    out_dir = str(docx_path.parent)
    pdf_path = docx_path.with_suffix(".pdf")
    if pdf_path.exists():
        pdf_path.unlink()

    env = os.environ.copy()
    env["SAL_USE_VCLPLUGIN"] = "svp"

    with tempfile.TemporaryDirectory(prefix="libreoffice-profile-") as profile_dir:
        result = subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={Path(profile_dir).resolve().as_uri()}",
                "--headless",
                "--nologo",
                "--nodefault",
                "--nolockcheck",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                out_dir,
                str(docx_path),
            ],
            env=env,
            capture_output=True,
            text=True,
            timeout=60,
        )

    deadline = time.monotonic() + 5
    while time.monotonic() < deadline and not pdf_path.exists():
        time.sleep(0.2)

    if pdf_path.exists():
        print(f"Interview prep PDF saved: {pdf_path}")
        if result.returncode != 0:
            detail = (result.stderr or result.stdout or "").strip()
            if detail:
                print(f"Note: LibreOffice non-zero exit after producing PDF: {detail}")
        return

    detail = (result.stderr or result.stdout or "").strip()
    if result.returncode == 0:
        print("Warning: PDF conversion ran but output file not found.")
    else:
        print(f"Warning: PDF conversion failed: {detail}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a personalized interview preparation guide.")
    parser.add_argument("output_dir", help="Job output directory (e.g. output/uber/senior-pm-web)")
    parser.add_argument(
        "--stage",
        default="General",
        choices=["General", "Recruiter Screen", "Phone Screen", "Onsite", "Final Round"],
        help="Interview stage to prepare for",
    )
    parser.add_argument(
        "--interviewer",
        action="append",
        default=[],
        dest="interviewers",
        help="Interviewer name (can be repeated)",
    )
    parser.add_argument("--notes", default="", help="Additional context or notes")
    parser.add_argument("--force", action="store_true", help="Regenerate even if guide already exists")
    args = parser.parse_args()

    out_dir = Path(args.output_dir)
    if not out_dir.is_absolute():
        out_dir = PROJECT_ROOT / out_dir

    if not out_dir.exists():
        print(f"Error: output directory does not exist: {out_dir}", file=sys.stderr)
        sys.exit(1)

    prep_dir = out_dir / "interview_prep"
    md_path = prep_dir / "interview_prep.md"

    # Check for existing guide
    if md_path.exists() and not args.force:
        print(f"Interview prep guide already exists: {md_path}")
        print("Use --force to regenerate.")
        sys.exit(0)

    # Step 1: Gather context
    _write_progress(prep_dir, "gathering_context", "Reading job description and candidate materials")
    print("Gathering context...")
    ctx = _gather_context(out_dir)
    company = ctx["company"]

    if not ctx["jd_raw"] and not ctx["jd_parsed"]:
        print(
            f"Error: no job description found in {out_dir / 'content'}. Run the pipeline first to generate JD content.",
            file=sys.stderr,
        )
        sys.exit(1)

    # Step 2: Build prompt
    _write_progress(prep_dir, "building_prompt", "Constructing prompt with all context")
    print("Building prompt...")
    full_prompt = _build_prompt(ctx, prep_dir, args.stage, args.interviewers, args.notes)

    # Step 3: Run configured provider
    provider = default_active_provider()
    _write_progress(prep_dir, "generating", f"{provider} provider generating prep guide for {company}")
    print(f"Generating interview prep guide for {ctx['jd_title']} at {company}...")
    print(f"This may take several minutes (up to 15 min) while {provider} researches and writes.")

    md_content = _run_provider(full_prompt, prep_dir)

    if not md_content:
        # Provider guard skipped interview prep (configured provider unavailable)
        _write_progress(prep_dir, "skipped", "Interview prep provider unavailable — skipped")
        print("Interview prep skipped (configured provider unavailable).")
        return

    _write_progress(prep_dir, "generating_documents", "Converting markdown to .docx and .pdf")

    # Step 4: Generate .docx
    print("Building .docx document...")
    docx_path = prep_dir / f"{CANDIDATE_NAME} Interview Prep - {company}.docx"
    _build_docx(md_content, docx_path)

    # Step 5: Generate .pdf
    print("Converting to PDF...")
    _convert_to_pdf(docx_path)

    _write_progress(prep_dir, "complete", "Interview prep guide ready")
    print(f"\nDone! Files in: {prep_dir}")


if __name__ == "__main__":
    main()
