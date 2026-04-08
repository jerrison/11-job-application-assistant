#!/usr/bin/env python3
"""Deterministic cover letter builder.

Takes a text file (or JSON with text field) and produces a .docx and .txt.

Usage:
    uv run scripts/build_cover_letter.py input.txt -o output/company/cover_letter.docx
    uv run scripts/build_cover_letter.py input.json -o output/company/cover_letter.docx

Input: plain text file with the cover letter body (starting with "Dear ..."),
       or a JSON file with a "text" field.

Formatting is fixed: Calibri 11pt, US Letter, 1in margins.
"""

import argparse
import json
import os
import shutil
import subprocess
import tempfile
import time
from pathlib import Path

from candidate_runtime import load_candidate_runtime_profile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from docx_text import sanitize_docx_text

FONT = "Calibri"
COLOR_BLACK = RGBColor(0x33, 0x33, 0x33)


def _run(para, text, size, bold=False, italic=False, color=COLOR_BLACK):
    run = para.add_run(sanitize_docx_text(text))
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>')
        rPr.insert(0, rFonts)
    return run


def build_cover_letter(text: str, output_path: str):
    doc = Document()
    candidate_profile = load_candidate_runtime_profile()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # Header: name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    _run(p, candidate_profile.full_name, Pt(14), bold=True)

    # Header: contact
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    _run(p, candidate_profile.contact_line(include_location=True), Pt(10), color=RGBColor(0x55, 0x55, 0x55))

    # Ensure greeting
    body = text.strip()
    import re

    if not re.match(r"(?i)^dear\s", body):
        body = "Dear Hiring Team,\n\n" + body

    # Ensure signoff
    if not re.search(r"\n(Best regards|Sincerely|Regards|Warm regards|Thank you),?\s*\n", body, re.IGNORECASE):
        body = body.rstrip() + f"\n\nBest regards,\n{candidate_profile.full_name}"

    # Body paragraphs
    paragraphs = [para.strip() for para in body.split("\n\n") if para.strip()]
    for para_text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        # Handle signoff line breaks (e.g. "Best regards,\nCandidate Name")
        lines = para_text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run("\n")
            _run(p, line, Pt(11))

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"Cover letter .docx saved: {output_path}")

    # Also save .txt (no header, just body for paste into forms)
    txt_path = output_path.replace(".docx", ".txt")
    with open(txt_path, "w") as f:
        f.write(text.strip())
    print(f"Cover letter .txt saved: {txt_path}")

    # Convert to PDF via LibreOffice
    _convert_to_pdf(output_path)


def _convert_to_pdf(docx_path: str):
    """Convert .docx to .pdf using LibreOffice."""
    soffice = _find_libreoffice()
    if not soffice:
        print("Warning: LibreOffice not found — skipping PDF conversion.")
        print("Install with: brew install --cask libreoffice")
        return

    out_dir = os.path.dirname(os.path.abspath(docx_path)) or "."
    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(out_dir, base + ".pdf")
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    env = os.environ.copy()
    # Force a true headless backend on macOS. The default osx VCL plugin can
    # initialize AppKit and abort even when `--headless` is present. This must
    # override inherited shell state such as SAL_USE_VCLPLUGIN=osx.
    env["SAL_USE_VCLPLUGIN"] = "svp"
    with tempfile.TemporaryDirectory(prefix="libreoffice-profile-") as profile_dir:
        result = subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={Path(profile_dir).resolve().as_uri()}",
                "--headless",
                "--nologo",
                "--nodefault",
                "--nolockcheck",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                out_dir,
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=60,
            env=env,
        )
    deadline = time.monotonic() + 5
    while time.monotonic() < deadline and not os.path.exists(pdf_path):
        time.sleep(0.2)
    if os.path.exists(pdf_path):
        print(f"Cover letter PDF saved: {pdf_path}")
        if result.returncode != 0:
            detail = (result.stderr or result.stdout or "").strip()
            if detail:
                print(f"Note: LibreOffice returned a non-zero exit code after producing the PDF: {detail}")
        return

    detail = (result.stderr or result.stdout or "").strip()
    if result.returncode == 0:
        print("Warning: PDF conversion ran but output file not found.")
    else:
        print(f"Warning: PDF conversion failed: {detail}")


def _find_libreoffice() -> str | None:
    """Find the LibreOffice soffice binary."""
    soffice = shutil.which("soffice")
    if soffice:
        return soffice
    mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.exists(mac_path):
        return mac_path
    return None


def main():
    parser = argparse.ArgumentParser(description="Build cover letter .docx from text")
    parser.add_argument("input", help="Path to text or JSON file")
    parser.add_argument("-o", "--output", default="output/cover_letter.docx", help="Output .docx path")
    args = parser.parse_args()

    with open(args.input) as f:
        content = f.read()

    # Try JSON first
    try:
        data = json.loads(content)
        text = data["text"]
    except (json.JSONDecodeError, KeyError):
        text = content

    build_cover_letter(text, args.output)


if __name__ == "__main__":
    main()
