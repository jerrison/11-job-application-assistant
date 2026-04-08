import importlib.util
import json
import sys
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent


def load_module(name: str, relative_path: str):
    path = PROJECT_ROOT / relative_path
    scripts_dir = str(path.parent)
    if scripts_dir not in sys.path:
        sys.path.insert(0, scripts_dir)
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


class ResumePolicyTests(unittest.TestCase):
    def test_draft_resume_uses_new_bullet_minimums(self):
        draft_resume = load_module("draft_resume", "scripts/draft_resume.py")
        self.assertEqual(draft_resume.SELECTION_RULES["moodys"], (7, 6))
        self.assertEqual(draft_resume.SELECTION_RULES["kyte"], (6, 5))
        self.assertEqual(draft_resume.SELECTION_RULES["tmobile"], (3, 3))
        self.assertEqual(draft_resume.SELECTION_RULES["lyft"], (1, 1))
        self.assertEqual(draft_resume.SELECTION_RULES["allstate"], (1, 1))

    def test_build_resume_uses_runtime_profile_for_contact_info(self):
        source = (PROJECT_ROOT / "scripts/build_resume.py").read_text()
        self.assertIn("load_candidate_runtime_profile", source)
        self.assertIn("_candidate_contact_for_input", source)
        self.assertNotIn("candidate@example.com", source)
        self.assertNotIn("linkedin.com/in/candidate/", source)

    def test_build_cover_letter_uses_runtime_profile_for_contact_info(self):
        source = (PROJECT_ROOT / "scripts/build_cover_letter.py").read_text()
        self.assertIn("load_candidate_runtime_profile", source)
        self.assertIn("candidate_profile.contact_line(include_location=True)", source)
        self.assertNotIn("candidate@example.com", source)
        self.assertNotIn("linkedin.com/in/candidate/", source)

    def test_docx_text_sanitizer_drops_xml_invalid_control_chars(self):
        docx_text = load_module("docx_text", "scripts/docx_text.py")

        sanitized = docx_text.sanitize_docx_text("Lead\x00PM\x08\x0b\t\n")

        self.assertEqual(sanitized, "LeadPM\t\n")

    def test_build_resume_run_sanitizes_xml_invalid_control_chars(self):
        build_resume = load_module("build_resume", "scripts/build_resume.py")
        paragraph = Document().add_paragraph()

        run = build_resume._run(paragraph, "Lead\x00PM\x0b", build_resume.Pt(11))

        self.assertEqual(run.text, "LeadPM")

    def test_build_cover_letter_run_sanitizes_xml_invalid_control_chars(self):
        build_cover_letter = load_module("build_cover_letter", "scripts/build_cover_letter.py")
        paragraph = Document().add_paragraph()

        run = build_cover_letter._run(paragraph, "Hiring\x00Team\x0b", build_cover_letter.Pt(11))

        self.assertEqual(run.text, "HiringTeam")

    def test_interview_prep_docx_run_sanitizes_xml_invalid_control_chars(self):
        generate_interview_prep = load_module("generate_interview_prep", "scripts/generate_interview_prep.py")
        paragraph = Document().add_paragraph()

        run = generate_interview_prep._docx_run(paragraph, "Impact\x00story\x0b", generate_interview_prep.Pt(11))

        self.assertEqual(run.text, "Impactstory")

    def test_build_resume_keeps_location_for_california_roles(self):
        build_resume = load_module("build_resume", "scripts/build_resume.py")

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            resume_path = root / "resume_content.json"
            jd_path = root / "jd_parsed.json"
            resume_path.write_text("{}", encoding="utf-8")
            jd_path.write_text(json.dumps({"location": "San Francisco, CA"}), encoding="utf-8")

            contact = build_resume._candidate_contact_for_input(resume_path)
            self.assertTrue(contact.startswith("San Francisco, CA"))

    def test_build_resume_omits_location_for_non_california_roles(self):
        build_resume = load_module("build_resume", "scripts/build_resume.py")

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            resume_path = root / "resume_content.json"
            jd_path = root / "jd_parsed.json"
            resume_path.write_text("{}", encoding="utf-8")
            jd_path.write_text(json.dumps({"location": "Denver, CO"}), encoding="utf-8")

            contact = build_resume._candidate_contact_for_input(resume_path)
            self.assertTrue(contact.startswith("candidate@example.com"))
            self.assertNotIn("San Francisco, CA", contact)

    def test_build_resume_keeps_location_when_any_listed_role_location_is_california(self):
        build_resume = load_module("build_resume", "scripts/build_resume.py")

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            resume_path = root / "resume_content.json"
            jd_path = root / "jd_parsed.json"
            application_page = root / "ashby_application_page.html"
            resume_path.write_text("{}", encoding="utf-8")
            jd_path.write_text(json.dumps({"location": "Denver"}), encoding="utf-8")
            application_page.write_text(
                '{"posting":{"locationName":"Denver","secondaryLocationNames":["Seattle","San Francisco","New York"]}}',
                encoding="utf-8",
            )

            contact = build_resume._candidate_contact_for_input(resume_path)
            self.assertTrue(contact.startswith("San Francisco, CA"))

    def test_build_resume_pdf_conversion_overrides_inherited_vcl_backend(self):
        build_resume = load_module("build_resume", "scripts/build_resume.py")

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            docx_path = root / "resume.docx"
            pdf_path = root / "resume.pdf"
            docx_path.write_bytes(b"docx")
            seen_env: dict[str, str] = {}

            def fake_run(cmd, **kwargs):
                nonlocal seen_env
                seen_env = dict(kwargs["env"])
                pdf_path.write_bytes(b"%PDF-1.4")
                return SimpleNamespace(returncode=0, stdout="", stderr="")

            with mock.patch.object(build_resume, "_find_libreoffice", return_value="/fake/soffice"):
                with mock.patch.object(build_resume.subprocess, "run", side_effect=fake_run):
                    with mock.patch.dict(build_resume.os.environ, {"SAL_USE_VCLPLUGIN": "osx"}, clear=False):
                        build_resume._convert_to_pdf(str(docx_path))

        self.assertEqual(seen_env["SAL_USE_VCLPLUGIN"], "svp")

    def test_build_cover_letter_pdf_conversion_overrides_inherited_vcl_backend(self):
        build_cover_letter = load_module("build_cover_letter", "scripts/build_cover_letter.py")

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            docx_path = root / "cover_letter.docx"
            pdf_path = root / "cover_letter.pdf"
            docx_path.write_bytes(b"docx")
            seen_env: dict[str, str] = {}

            def fake_run(cmd, **kwargs):
                nonlocal seen_env
                seen_env = dict(kwargs["env"])
                pdf_path.write_bytes(b"%PDF-1.4")
                return SimpleNamespace(returncode=0, stdout="", stderr="")

            with mock.patch.object(build_cover_letter, "_find_libreoffice", return_value="/fake/soffice"):
                with mock.patch.object(build_cover_letter.subprocess, "run", side_effect=fake_run):
                    with mock.patch.dict(build_cover_letter.os.environ, {"SAL_USE_VCLPLUGIN": "osx"}, clear=False):
                        build_cover_letter._convert_to_pdf(str(docx_path))

        self.assertEqual(seen_env["SAL_USE_VCLPLUGIN"], "svp")

    def test_enforce_resume_policy_tops_up_required_positions(self):
        enforce_resume_policy = load_module("enforce_resume_policy", "scripts/enforce_resume_policy.py")

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            resume_path = root / "resume_content.json"
            ranked_path = root / "ranked_bullets.json"

            resume_path.write_text(
                json.dumps(
                    {
                        "tagline": "Test",
                        "summary": None,
                        "positions": {
                            "moodys": [
                                {"bold": "Existing Moody's bullet 1, ", "text": "text"},
                                {"bold": "Existing Moody's bullet 2, ", "text": "text"},
                                {"bold": "Existing Moody's bullet 3, ", "text": "text"},
                                {"bold": "Existing Moody's bullet 4, ", "text": "text"},
                            ],
                            "kyte": [
                                {"bold": "Existing Kyte bullet 1, ", "text": "text"},
                                {"bold": "Existing Kyte bullet 2, ", "text": "text"},
                                {"bold": "Existing Kyte bullet 3, ", "text": "text"},
                            ],
                            "tmobile": [
                                {"bold": "Existing T-Mobile bullet 1, ", "text": "text"},
                                {"bold": "Existing T-Mobile bullet 2, ", "text": "text"},
                            ],
                            "lyft": [],
                            "allstate": [],
                        },
                        "page_break_before": None,
                    }
                )
            )

            ranked_path.write_text(
                json.dumps(
                    {
                        "positions": {
                            "moodys": [
                                {
                                    "score": 0.9,
                                    "bullet": "Launched catastrophe workflow automation program, improving underwriter turnaround for enterprise carriers.",
                                },
                                {
                                    "score": 0.8,
                                    "bullet": "Secured enterprise insurance growth through pricing-platform roadmap decisions and executive partnership.",
                                },
                            ],
                            "kyte": [
                                {
                                    "score": 0.9,
                                    "bullet": "Built fraud-loss model for claims exposure, improving ML-driven risk decisions.",
                                },
                                {
                                    "score": 0.8,
                                    "bullet": "Established experimentation platform for marketplace pricing and product analytics.",
                                },
                            ],
                            "tmobile": [
                                {
                                    "score": 0.9,
                                    "bullet": "Enabled IoT OEM co-selling channel with a partner onboarding portal and analytics workspace.",
                                },
                            ],
                            "lyft": [
                                {
                                    "score": 0.9,
                                    "bullet": "Reduced company-wide insurance cost through predictive risk modeling in a regulated environment.",
                                },
                            ],
                            "allstate": [
                                {
                                    "score": 0.9,
                                    "bullet": "Built internal pricing tooling that reduced variance and saved quarterly analyst time.",
                                },
                            ],
                        }
                    }
                )
            )

            changed, actions = enforce_resume_policy.enforce_policies(resume_path)

            self.assertTrue(changed)
            self.assertEqual(len(actions), 7)

            updated = json.loads(resume_path.read_text())
            self.assertEqual(len(updated["positions"]["moodys"]), 6)
            self.assertEqual(len(updated["positions"]["kyte"]), 5)
            self.assertEqual(len(updated["positions"]["tmobile"]), 3)
            self.assertEqual(len(updated["positions"]["lyft"]), 1)
            self.assertEqual(len(updated["positions"]["allstate"]), 1)


if __name__ == "__main__":
    unittest.main()
