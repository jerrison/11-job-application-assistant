import importlib.util
import io
import unittest
from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas

PROJECT_ROOT = Path(__file__).resolve().parent.parent


def load_module(name: str, relative_path: str):
    path = PROJECT_ROOT / relative_path
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


class MaterialIngestTests(unittest.TestCase):
    def test_import_text_payload_returns_plain_text(self):
        material_ingest = load_module("material_ingest", "scripts/material_ingest.py")

        text = material_ingest.import_material_content(
            text="Plain text resume\nwith bullets\n",
        )

        self.assertEqual(text, "Plain text resume\nwith bullets\n")

    def test_import_docx_payload_extracts_paragraph_text(self):
        material_ingest = load_module("material_ingest", "scripts/material_ingest.py")

        buffer = io.BytesIO()
        document = Document()
        document.add_paragraph("Imported from DOCX")
        document.add_paragraph("Second paragraph")
        document.save(buffer)

        text = material_ingest.import_material_content(
            file_name="resume.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content_bytes=buffer.getvalue(),
        )

        self.assertIn("Imported from DOCX", text)
        self.assertIn("Second paragraph", text)

    def test_import_pdf_payload_extracts_page_text(self):
        material_ingest = load_module("material_ingest", "scripts/material_ingest.py")

        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer)
        pdf.drawString(72, 720, "Imported from PDF")
        pdf.save()

        text = material_ingest.import_material_content(
            file_name="resume.pdf",
            content_type="application/pdf",
            content_bytes=buffer.getvalue(),
        )

        self.assertIn("Imported from PDF", text)

    def test_import_public_google_doc_url_uses_export_text_endpoint(self):
        material_ingest = load_module("material_ingest", "scripts/material_ingest.py")
        seen_urls: list[str] = []

        def fake_fetch(url: str):
            seen_urls.append(url)
            return material_ingest.FetchResult(
                url=url,
                content_type="text/plain; charset=utf-8",
                content=b"Imported from URL\n",
            )

        text = material_ingest.import_material_content(
            source_url="https://docs.google.com/document/d/abc123/edit?tab=t.0",
            fetcher=fake_fetch,
        )

        self.assertEqual(text, "Imported from URL\n")
        self.assertEqual(
            seen_urls,
            ["https://docs.google.com/document/d/abc123/export?format=txt"],
        )


if __name__ == "__main__":
    unittest.main()
